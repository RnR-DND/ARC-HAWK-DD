"""DOCX connector — uses python-docx for Word document extraction."""

import os
import logging
from typing import Any, Generator

from hawk_scanner.commands.base import BaseConnector, FieldRecord, validate_magic_bytes
from hawk_scanner.commands.base import register_connector

logger = logging.getLogger(__name__)


@register_connector("docx")
class DocxConnector(BaseConnector):
    connector_type = "docx"

    def __init__(self, config: dict):
        super().__init__(config)
        self._file_path = config.get("path", "")

    def connect(self) -> None:
        if not os.path.exists(self._file_path):
            raise FileNotFoundError(f"File not found: {self._file_path}")

        if not validate_magic_bytes(self._file_path, "docx"):
            raise ValueError(f"Invalid DOCX (magic bytes mismatch): {self._file_path}")

        self._connected = True
        self.logger.info("DOCX connector ready for %s", self._file_path)

    def discover(self) -> list[dict[str, Any]]:
        from docx import Document

        doc = Document(self._file_path)
        num_paragraphs = len(doc.paragraphs)
        num_tables = len(doc.tables)

        return [{
            "name": os.path.basename(self._file_path),
            "schema": os.path.dirname(self._file_path),
            "table": os.path.basename(self._file_path),
            "type": "docx",
            "num_paragraphs": num_paragraphs,
            "num_tables": num_tables,
            "columns": [
                {"name": "paragraph_text", "data_type": "text", "nullable": True},
                {"name": "table_cell", "data_type": "text", "nullable": True},
            ],
        }]

    def extract(self, target: str, sample_size: int = 1000) -> Generator[FieldRecord, None, None]:
        from docx import Document

        source = self._make_source()
        doc = Document(self._file_path)
        count = 0

        # Extract paragraphs
        for i, para in enumerate(doc.paragraphs):
            if count >= sample_size:
                break
            text = para.text.strip()
            if text:
                yield FieldRecord(
                    field=f"paragraph_{i + 1}",
                    value=text,
                    source=source,
                    field_path=f"{os.path.basename(self._file_path)}.paragraph_{i + 1}",
                    metadata={
                        "connector": self.connector_type,
                        "file": self._file_path,
                        "style": para.style.name if para.style else "Normal",
                    },
                )
                count += 1

        # Extract table cells
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if count >= sample_size:
                        return
                    text = cell.text.strip()
                    if text:
                        yield FieldRecord(
                            field=f"table_{table_idx + 1}_row_{row_idx + 1}_col_{cell_idx + 1}",
                            value=text,
                            source=source,
                            field_path=(
                                f"{os.path.basename(self._file_path)}"
                                f".table_{table_idx + 1}.row_{row_idx + 1}.col_{cell_idx + 1}"
                            ),
                            metadata={
                                "connector": self.connector_type,
                                "file": self._file_path,
                                "content_type": "table_cell",
                            },
                        )
                        count += 1

    def close(self) -> None:
        self._connected = False

    def _make_source(self) -> str:
        return f"docx://{self._file_path}"
